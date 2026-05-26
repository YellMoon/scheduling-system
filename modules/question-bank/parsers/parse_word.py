#!/usr/bin/env python3
"""Word question parser for the question-bank import pipeline.

Usage: python parse_word.py <file_path> <source_type> [knowledge_tree_path]
source_type: lecture | exam | auto
"""

import io
import base64
import hashlib
import html
import json
import mimetypes
import os
import re
import subprocess
import sys
import tempfile
import uuid
import zipfile
import xml.etree.ElementTree as ET

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8")


QUESTION_RE = re.compile(r"^(\d+)[\.\u3001\uff0e]\s*(.*)")
OPTION_RE = re.compile(r"^([A-G])[\.\uff0e]\s*(.*)", re.I)
SUB_QUESTION_RE = re.compile(r"^(?:[\(\uff08](\d+)[\)\uff09]|([\u2460\u2461\u2462\u2463\u2464\u2465\u2466\u2467\u2468\u2469]))\s*(.*)")
SYMBOL_FONT_MAP = str.maketrans({
    "a": "α", "b": "β", "c": "χ", "d": "δ", "e": "ε", "f": "φ", "g": "γ", "h": "η",
    "i": "ι", "j": "ϕ", "k": "κ", "l": "λ", "m": "μ", "n": "ν", "o": "ο", "p": "π",
    "q": "θ", "r": "ρ", "s": "σ", "t": "τ", "u": "υ", "v": "ϖ", "w": "ω", "x": "ξ",
    "y": "ψ", "z": "ζ", "A": "Α", "B": "Β", "C": "Χ", "D": "Δ", "E": "Ε", "F": "Φ",
    "G": "Γ", "H": "Η", "I": "Ι", "J": "ϑ", "K": "Κ", "L": "Λ", "M": "Μ", "N": "Ν",
    "O": "Ο", "P": "Π", "Q": "Θ", "R": "Ρ", "S": "Σ", "T": "Τ", "U": "Υ", "V": "ς",
    "W": "Ω", "X": "Ξ", "Y": "Ψ", "Z": "Ζ",
})
EXAM_SECTION_RE = re.compile(r"^[\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341]+[\u3001\uff0e\.\s]*(\u5355\u9009\u9898|\u591a\u9009\u9898|\u9009\u62e9\u9898|\u5b9e\u9a8c\u9898|\u89e3\u7b54\u9898|\u7efc\u5408\u9898|\u586b\u7a7a\u9898)")
ANSWER_TITLE_RE = re.compile(r"^(?:\u300a.*?\u300b\s*)?(?:\u53c2\u8003\u7b54\u6848|\u7b54\u6848\u4e0e\u89e3\u6790|\u7b54\u6848\u53ca\u89e3\u6790|\u7b54\u6848)")
ANALYSIS_MARK_RE = re.compile(r"\u3010(?!\u7b54\u6848)[^】]+\u3011")


def extract_question_number(text):
    match = QUESTION_RE.match(text)
    if match:
        return int(match.group(1)), match.group(2).strip()
    return None, text


def read_numbering_definitions(file_path):
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    try:
        with zipfile.ZipFile(file_path, "r") as archive:
            if "word/numbering.xml" not in archive.namelist():
                return {}
            root = ET.fromstring(archive.read("word/numbering.xml"))
            num_to_abstract = {}
            for num in root.findall("w:num", namespace):
                num_id = num.attrib.get(f"{{{namespace['w']}}}numId")
                abstract = num.find("w:abstractNumId", namespace)
                if num_id and abstract is not None:
                    num_to_abstract[num_id] = abstract.attrib.get(f"{{{namespace['w']}}}val")

            abstract_levels = {}
            for abstract in root.findall("w:abstractNum", namespace):
                abstract_id = abstract.attrib.get(f"{{{namespace['w']}}}abstractNumId")
                levels = {}
                for level in abstract.findall("w:lvl", namespace):
                    ilvl = level.attrib.get(f"{{{namespace['w']}}}ilvl", "0")
                    lvl_text = level.find("w:lvlText", namespace)
                    num_fmt = level.find("w:numFmt", namespace)
                    start = level.find("w:start", namespace)
                    levels[ilvl] = {
                        "text": lvl_text.attrib.get(f"{{{namespace['w']}}}val", "") if lvl_text is not None else "",
                        "format": num_fmt.attrib.get(f"{{{namespace['w']}}}val", "") if num_fmt is not None else "",
                        "start": int(start.attrib.get(f"{{{namespace['w']}}}val", "1")) if start is not None else 1,
                    }
                abstract_levels[abstract_id] = levels

            definitions = {}
            for num_id, abstract_id in num_to_abstract.items():
                definitions[num_id] = abstract_levels.get(abstract_id, {})
            return definitions
    except Exception:
        return {}


def paragraph_numbering(paragraph):
    try:
        num_pr = paragraph._p.pPr.numPr if paragraph._p.pPr is not None else None
        if num_pr is None or num_pr.numId is None:
            return None
        num_id = str(num_pr.numId.val)
        ilvl = str(num_pr.ilvl.val) if num_pr.ilvl is not None else "0"
        return num_id, ilvl
    except Exception:
        return None


def _rels_for_part(archive, rel_path, base_prefix="word/"):
    rels = {}
    if rel_path not in archive.namelist():
        return rels
    root = ET.fromstring(archive.read(rel_path))
    for rel in root:
        rid = rel.attrib.get("Id")
        target = rel.attrib.get("Target", "")
        rel_type = rel.attrib.get("Type", "")
        if not rid or not target:
            continue
        if target.startswith("/"):
            full = target.lstrip("/")
        elif target.startswith("word/"):
            full = target
        else:
            full = os.path.normpath(base_prefix + target).replace("\\", "/")
        rels[rid] = {"target": full, "type": rel_type}
    return rels


def read_style_names_from_archive(archive):
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    if "word/styles.xml" not in archive.namelist():
        return {}
    try:
        root = ET.fromstring(archive.read("word/styles.xml"))
        styles = {}
        for style in root.findall(".//w:style", ns):
            style_id = style.attrib.get("{%s}styleId" % ns["w"], "")
            name_node = style.find("./w:name", ns)
            name = name_node.attrib.get("{%s}val" % ns["w"], "") if name_node is not None else ""
            if style_id:
                styles[style_id] = name or style_id
        return styles
    except Exception:
        return {}


def read_docx_comments(file_path):
    namespace = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "v": "urn:schemas-microsoft-com:vml",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "o": "urn:schemas-microsoft-com:office:office",
        "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    }
    try:
        with zipfile.ZipFile(file_path, "r") as archive:
            if "word/comments.xml" not in archive.namelist():
                return {}
            root = ET.fromstring(archive.read("word/comments.xml"))
            rels = _rels_for_part(archive, "word/_rels/comments.xml.rels")
            comments = {}
            for comment in root.findall(".//w:comment", namespace):
                comment_id = comment.attrib.get(f"{{{namespace['w']}}}id")
                texts = []
                for paragraph in comment.findall(".//w:p", namespace):
                    inline_assets = []
                    text = _paragraph_text_with_markup(paragraph, namespace, archive, rels, bool(paragraph.findall(".//o:OLEObject", namespace)), inline_assets)
                    if text:
                        texts.append(text)
                if comment_id:
                    comments[comment_id] = "\n".join(texts).strip()
            return comments
    except Exception:
        return {}


def read_docx_comment_assets(file_path):
    namespace = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "v": "urn:schemas-microsoft-com:vml",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "o": "urn:schemas-microsoft-com:office:office",
        "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    }
    assets = []
    seen = set()
    try:
        with zipfile.ZipFile(file_path, "r") as archive:
            if "word/comments.xml" not in archive.namelist():
                return []
            root = ET.fromstring(archive.read("word/comments.xml"))
            rels = _rels_for_part(archive, "word/_rels/comments.xml.rels")
            for paragraph in root.findall(".//w:p", namespace):
                inline_assets = []
                _paragraph_text_with_markup(paragraph, namespace, archive, rels, bool(paragraph.findall(".//o:OLEObject", namespace)), inline_assets)
                for asset in inline_assets:
                    key = asset.get("content_hash")
                    if key and key not in seen:
                        assets.append(asset)
                        seen.add(key)
    except Exception:
        return []
    return assets


def read_paragraph_comment_refs(file_path):
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    try:
        with zipfile.ZipFile(file_path, "r") as archive:
            if "word/document.xml" not in archive.namelist():
                return []
            root = ET.fromstring(archive.read("word/document.xml"))
            refs_by_paragraph = []
            for paragraph in root.findall(".//w:p", namespace):
                refs = []
                for node in paragraph.findall(".//w:commentReference", namespace) + paragraph.findall(".//w:commentRangeStart", namespace) + paragraph.findall(".//w:commentRangeEnd", namespace):
                    comment_id = node.attrib.get(f"{{{namespace['w']}}}id")
                    if comment_id and comment_id not in refs:
                        refs.append(comment_id)
                refs_by_paragraph.append(refs)
            return refs_by_paragraph
    except Exception:
        return []


def extract_numbered_items(doc, file_path):
    definitions = read_numbering_definitions(file_path)
    comments = read_docx_comments(file_path)
    counters = {}
    items = []
    xml_items = extract_numbered_items_from_xml(file_path, definitions, comments, counters)
    if xml_items:
        return xml_items

    comment_refs = read_paragraph_comment_refs(file_path)
    for index, paragraph in enumerate(doc.paragraphs):
        text = clean_word_text(paragraph.text)
        numbering = paragraph_numbering(paragraph)
        number_label = ""
        number_kind = ""
        if numbering:
            num_id, ilvl = numbering
            level = definitions.get(num_id, {}).get(ilvl, {})
            lvl_text = level.get("text", "")
            if level.get("format") == "decimal" and "%1" in lvl_text and ilvl == "0":
                key = (num_id, ilvl)
                counters[key] = counters.get(key, level.get("start", 1) - 1) + 1
                number_label = lvl_text.replace("%1", str(counters[key]))
                if lvl_text.startswith("例"):
                    number_kind = "example"
                elif lvl_text in ("%1", "%1.", "%1．"):
                    number_kind = "practice"
        paragraph_comments = [
            comments[comment_id]
            for comment_id in (comment_refs[index] if index < len(comment_refs) else [])
            if comments.get(comment_id)
        ]
        style_name = getattr(getattr(paragraph, "style", None), "name", "") or ""
        items.append({
            "text": text,
            "number_label": number_label,
            "number_kind": number_kind,
            "comments": paragraph_comments,
            "is_heading": style_name.lower().startswith("heading") or "\u6807\u9898" in style_name,
        })
    return items


def normalize_operator_symbols(text):
    if not text:
        return ""
    protected_tags = []

    def protect_tag(match):
        token = "@@HTML_TAG_%d@@" % len(protected_tags)
        protected_tags.append(match.group(0))
        return token

    value = re.sub(r"<[^>]+>", protect_tag, text)
    value = value.replace("\u2212", "\uff0d").replace("+", "\uff0b").replace("-", "\uff0d")
    return re.sub(
        r"@@HTML_TAG_(\d+)@@",
        lambda m: protected_tags[int(m.group(1))] if int(m.group(1)) < len(protected_tags) else "",
        value,
    )


def clean_word_text(text):
    return normalize_operator_symbols(re.sub(r"\v+", "", (text or "").replace("\f", "\n")).strip())


def visible_space_text(text):
    if not text:
        return ""
    return re.sub(r" {2,}", lambda m: "&nbsp;" * len(m.group(0)), text.replace("\u3000", " "))


def _xml_bytes(node):
    try:
        return ET.tostring(node, encoding="unicode")
    except Exception:
        return ""

def _local_name(node):
    return node.tag.rsplit("}", 1)[-1]

def _first_child(node, name):
    for child in list(node):
        if _local_name(child) == name:
            return child
    return None


def _omml_val(node, child_name, default=""):
    child = _first_child(node, child_name) if node is not None else None
    if child is None:
        return default
    for key, value in child.attrib.items():
        if key.endswith("}val") or key == "val":
            return value
    return default


def _omml_fraction_type(node):
    fpr = _first_child(node, "fPr")
    return _omml_val(fpr, "type", "bar") if fpr is not None else "bar"

def _format_math_token(text):
    if not text:
        return ""
    if re.fullmatch(r"[A-Za-z]", text):
        return "<i>%s</i>" % text
    return re.sub(r"([A-Za-z])", r"<i>\1</i>", text)


UNIT_SYMBOLS = [
    "kg", "mol", "cd", "rad", "sr", "Hz", "Pa", "J", "W", "C", "V", "F", "Ω", "S", "Wb", "T", "H", "lm", "lx", "Bq", "Gy", "Sv", "kat",
    "m", "s", "A", "K", "N", "Pa", "J", "W", "C", "V", "F", "T", "H", "L", "eV", "MeV", "GeV", "min", "h",
]
UNIT_PREFIXES = ["Y", "Z", "E", "P", "T", "G", "M", "k", "h", "da", "d", "c", "m", "μ", "u", "n", "p", "f", "a", "z", "y"]


def _unit_tokens():
    tokens = set(UNIT_SYMBOLS)
    for prefix in UNIT_PREFIXES:
        for unit in UNIT_SYMBOLS:
            if unit in ("kg", "min", "mol", "rad", "sr"):
                continue
            tokens.add(prefix + unit)
    return sorted(tokens, key=len, reverse=True)


UNIT_PATTERN = r"(?:%s)(?:\s*(?:[·⋅*/\/]\s*|\^?-?\d+|<sup>-?\d+</sup>|\s+)(?:%s))*" % (
    "|".join(re.escape(token) for token in _unit_tokens()),
    "|".join(re.escape(token) for token in _unit_tokens()),
)


def normalize_physics_markup(text):
    if not text:
        return ""
    protected_images = []

    def protect_image(match):
        token = "@@QUESTION_IMAGE_%d@@" % len(protected_images)
        protected_images.append(match.group(0))
        return token

    text = re.sub(r"<img\b[^>]*>", protect_image, text, flags=re.I)
    text = re.sub(r"<i>k</i><i>g</i>", "kg", text)
    text = re.sub(r"<i>(N|Pa|J|W|C|V|F|T|H|A|K|Hz|Ω)</i>(?=(?:<sup>|[·⋅.*/\/]))", r"\1", text)
    text = re.sub(r"(?<=[·⋅.*/\/])<i>(m|s|g|A|K|mol|cd|rad|Hz|N|Pa|J|W|C|V|F|T|H|Ω)</i>", r"\1", text)
    text = re.sub(r"<i>(m|s|g|A|K|mol|cd|rad|Hz|N|Pa|J|W|C|V|F|T|H|Ω)</i>(?=<sup>-?\d+</sup>)", r"\1", text)
    def restore_unit(match):
        prefix = match.group(1)
        body = match.group(2)
        return prefix + re.sub(r"</?i>", "", body)

    italic_unit_token = r"(?:<i>[A-Za-zμΩ]+</i>|[·⋅*/\/\s]|\^?-?\d+|<sup>-?\d+</sup>)+"
    text = re.sub(r"(\d(?:[\d.,×xX\-+]*\d)?\s*)(" + italic_unit_token + r")(?=\s|[\u4e00-\u9fff]|[,，。；;、）)]|$)", restore_unit, text)
    text = re.sub(r"(?<=\d)\s*(%s)" % UNIT_PATTERN, lambda m: " " + m.group(1), text)
    text = re.sub(r"@@QUESTION_IMAGE_(\d+)@@", lambda m: protected_images[int(m.group(1))] if int(m.group(1)) < len(protected_images) else "", text)
    return text

def _plain_math_text(node):
    if node is None:
        return ""
    tag = _local_name(node)
    if tag == "t":
        return node.text or ""
    return "".join(_plain_math_text(child) for child in list(node))

def _math_text(node):
    if node is None:
        return ""
    tag = _local_name(node)
    if tag == "t":
        return _format_math_token(node.text or "")
    if tag == "r":
        plain_style = False
        rpr = _first_child(node, "rPr")
        if rpr is not None:
            plain_style = any(_local_name(child) in ("nor", "lit") for child in list(rpr))
            plain_style = plain_style or any(_local_name(child) == "sty" and child.attrib.get("{http://schemas.openxmlformats.org/officeDocument/2006/math}val") in ("p", "plain") for child in list(rpr))
        if plain_style:
            return _plain_math_text(node)
        return "".join(_math_text(child) for child in list(node) if _local_name(child) != "rPr")
    if tag == "fName":
        return _plain_math_text(node)
    if tag in ("oMath", "oMathPara", "e", "num", "den", "deg", "sub", "sup"):
        return "".join(_math_text(child) for child in list(node))
    if tag == "f":
        num = _math_text(_first_child(node, "num"))
        den = _math_text(_first_child(node, "den"))
        if _omml_fraction_type(node) in ("lin", "skw"):
            return "%s/%s" % (num, den)
        return '<span class="omml-frac"><span class="omml-frac-num">%s</span><span class="omml-frac-den">%s</span></span>' % (num, den)
    if tag == "sSub":
        base = _math_text(_first_child(node, "e"))
        sub = _math_text(_first_child(node, "sub"))
        return "%s<sub>%s</sub>" % (base, sub)
    if tag == "sSup":
        base = _math_text(_first_child(node, "e"))
        sup = _math_text(_first_child(node, "sup"))
        return "%s<sup>%s</sup>" % (base, sup)
    if tag == "sSubSup":
        base = _math_text(_first_child(node, "e"))
        sub = _math_text(_first_child(node, "sub"))
        sup = _math_text(_first_child(node, "sup"))
        if not base:
            return '<span class="nuclear-left"><sup>%s</sup><sub>%s</sub></span>' % (sup, sub)
        return "%s<sub>%s</sub><sup>%s</sup>" % (base, sub, sup)
    if tag == "rad":
        deg = _math_text(_first_child(node, "deg"))
        body = _math_text(_first_child(node, "e"))
        latex = _math_latex(node)
        if latex:
            return _legacy_latex_span(latex)
        rad_class = "omml-rad"
        if not deg:
            return '<span class="%s"><span class="omml-rad-sign">&#8730;</span><span class="omml-rad-body">%s</span></span>' % (rad_class, body)
        return '<span class="%s"><span class="omml-rad-index">%s</span><span class="omml-rad-sign">&#8730;</span><span class="omml-rad-body">%s</span></span>' % (rad_class, deg, body)
    if tag == "nary":
        symbol = "".join(child.attrib.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "") for child in node.iter() if _local_name(child) == "chr") or "∑"
        sub = _math_text(_first_child(node, "sub"))
        sup = _math_text(_first_child(node, "sup"))
        body = _math_text(_first_child(node, "e"))
        return "%s%s%s%s" % (symbol, "<sub>%s</sub>" % sub if sub else "", "<sup>%s</sup>" % sup if sup else "", body)
    if tag == "d":
        body = _math_text(_first_child(node, "e"))
        return "(%s)" % body
    if tag == "bar":
        return r"$\overline{%s}$" % _math_text(_first_child(node, "e"))
    if tag == "groupChr":
        return _math_text(_first_child(node, "e"))
    if tag == "func":
        return "%s(%s)" % (_plain_math_text(_first_child(node, "fName")), _math_text(_first_child(node, "e")))
    return "".join(_math_text(child) for child in list(node))


def _legacy_latex_span(latex):
    if not latex:
        return ""
    return '<span class="legacy-latex" data-latex="%s"></span>' % html.escape(latex, quote=True)


def _math_latex(node):
    if node is None:
        return ""
    tag = _local_name(node)
    if tag == "t":
        return node.text or ""
    if tag == "r":
        plain_style = False
        rpr = _first_child(node, "rPr")
        if rpr is not None:
            plain_style = any(_local_name(child) in ("nor", "lit") for child in list(rpr))
            plain_style = plain_style or any(_local_name(child) == "sty" and child.attrib.get("{http://schemas.openxmlformats.org/officeDocument/2006/math}val") in ("p", "plain") for child in list(rpr))
        if plain_style:
            return _plain_math_text(node)
        return "".join(_math_latex(child) for child in list(node) if _local_name(child) != "rPr")
    if tag == "fName":
        return _plain_math_text(node)
    if tag in ("oMath", "oMathPara", "e", "num", "den", "deg", "sub", "sup"):
        return "".join(_math_latex(child) for child in list(node))
    if tag == "f":
        num = _math_latex(_first_child(node, "num"))
        den = _math_latex(_first_child(node, "den"))
        if _omml_fraction_type(node) in ("lin", "skw"):
            return "%s/%s" % (num, den)
        return r"\frac{%s}{%s}" % (num, den)
    if tag == "sSub":
        base = _math_latex(_first_child(node, "e"))
        sub = _math_latex(_first_child(node, "sub"))
        return "%s_{%s}" % (base, sub)
    if tag == "sSup":
        base = _math_latex(_first_child(node, "e"))
        sup = _math_latex(_first_child(node, "sup"))
        return "%s^{%s}" % (base, sup)
    if tag == "sSubSup":
        base = _math_latex(_first_child(node, "e"))
        sub = _math_latex(_first_child(node, "sub"))
        sup = _math_latex(_first_child(node, "sup"))
        return "%s_{%s}^{%s}" % (base, sub, sup) if base else "{}_{%s}^{%s}" % (sub, sup)
    if tag == "rad":
        deg = _math_latex(_first_child(node, "deg"))
        body = _math_latex(_first_child(node, "e"))
        return r"\sqrt[%s]{%s}" % (deg, body) if deg else r"\sqrt{%s}" % body
    if tag == "nary":
        symbol = "".join(child.attrib.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "") for child in node.iter() if _local_name(child) == "chr")
        command = {"∑": r"\sum", "Σ": r"\sum", "∫": r"\int", "∏": r"\prod"}.get(symbol, symbol or r"\sum")
        sub = _math_latex(_first_child(node, "sub"))
        sup = _math_latex(_first_child(node, "sup"))
        body = _math_latex(_first_child(node, "e"))
        limits = ("%s%s" % ("_{%s}" % sub if sub else "", "^{%s}" % sup if sup else ""))
        return "%s%s %s" % (command, limits, body)
    if tag == "d":
        return r"\left(%s\right)" % _math_latex(_first_child(node, "e"))
    if tag == "bar":
        return r"\overline{%s}" % _math_latex(_first_child(node, "e"))
    if tag == "groupChr":
        return _math_latex(_first_child(node, "e"))
    if tag == "func":
        return "%s(%s)" % (_plain_math_text(_first_child(node, "fName")), _math_latex(_first_child(node, "e")))
    return "".join(_math_latex(child) for child in list(node))

def _asset_from_run_rel(archive, rels, rid, has_ole_object=False):
    if archive is None or not rid:
        return None
    rel = rels.get(rid or "") if rels else None
    return _asset_from_part(archive, rel.get("target") if rel else "", "formula_preview" if has_ole_object else "image", rid, rel.get("type") if rel else "")


EMU_PER_PIXEL = 9525


def _display_size_from_container(container, ns):
    extent = container.find(".//wp:extent", ns) if "wp" in ns else None
    if extent is not None:
        try:
            width = round(int(extent.attrib.get("cx", "0")) / EMU_PER_PIXEL)
            height = round(int(extent.attrib.get("cy", "0")) / EMU_PER_PIXEL)
            if width > 0 and height > 0:
                return {"display_width": width, "display_height": height}
        except Exception:
            pass
    shape = container.find(".//v:shape", ns) if "v" in ns else None
    style = shape.attrib.get("style", "") if shape is not None else ""
    if style:
        width_match = re.search(r"width\s*:\s*([\d.]+)pt", style)
        height_match = re.search(r"height\s*:\s*([\d.]+)pt", style)
        if width_match and height_match:
            return {
                "display_width": round(float(width_match.group(1)) * 96 / 72),
                "display_height": round(float(height_match.group(1)) * 96 / 72),
            }
    return {}


def _image_tag(asset):
    src = "question-asset://%s" % asset.get("content_hash", "") if asset.get("content_hash") else asset.get("data_url", "")
    alt = asset.get("file_name", "image")
    width = asset.get("display_width")
    height = asset.get("display_height")
    size_attrs = ""
    if width and height:
        size_attrs = ' width="%s" height="%s" style="width:%spx;height:%spx;"' % (width, height, width, height)
    return '<img src="%s" alt="%s"%s />' % (src, alt, size_attrs)


def _paragraph_text_with_markup(paragraph, ns, archive=None, rels=None, has_ole_object=False, inline_assets=None):
    parts = []
    inline_assets = inline_assets if inline_assets is not None else []

    def has_symbol_font(run):
        fonts = run.findall("./w:rPr/w:rFonts", ns)
        for font in fonts:
            values = [
                font.attrib.get("{%s}%s" % (ns.get("w", ""), key), "")
                for key in ("ascii", "hAnsi", "eastAsia", "cs")
            ]
            if any(str(value).lower() == "symbol" for value in values):
                return True
        return False

    def direct_text_outside_graphics(run, tag_name):
        try:
            safe_run = ET.fromstring(_xml_bytes(run))
        except Exception:
            safe_run = run

        def remove_graphic_children(node):
            for child in list(node):
                if _local_name(child) in ("drawing", "pict", "object", "oMath", "oMathPara"):
                    node.remove(child)
                else:
                    remove_graphic_children(child)

        remove_graphic_children(safe_run)
        values = []
        for node in safe_run.iter():
            local = _local_name(node)
            if local == tag_name:
                values.append(node.text or "")
            elif local == "tab":
                values.append("\t")
            elif local == "br":
                br_type = node.attrib.get("{%s}type" % ns.get("w", ""))
                values.append("\f" if br_type == "page" else "\n")
        return "".join(values)

    for child in list(paragraph):
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "r":
            if archive is not None and all(key in ns for key in ("a", "v", "r")):
                for blip in child.findall(".//a:blip", ns):
                    rid = blip.attrib.get("{%s}embed" % ns.get("r", "")) or blip.attrib.get("{%s}link" % ns.get("r", ""))
                    asset = _asset_from_run_rel(archive, rels, rid, has_ole_object)
                    if asset:
                        asset.update(_display_size_from_container(child, ns))
                        inline_assets.append(asset)
                        parts.append(_image_tag(asset))
                for image in child.findall(".//v:imagedata", ns):
                    rid = image.attrib.get("{%s}id" % ns.get("r", ""))
                    asset = _asset_from_run_rel(archive, rels, rid, has_ole_object)
                    if asset:
                        asset.update(_display_size_from_container(child, ns))
                        inline_assets.append(asset)
                        parts.append(_image_tag(asset))
            for math_node in child.findall(".//m:oMath", ns) + child.findall(".//m:oMathPara", ns):
                math_latex = _math_latex(math_node).strip()
                if math_latex:
                    parts.append(_legacy_latex_span(math_latex))
            text = direct_text_outside_graphics(child, "t")
            if not text:
                continue
            if has_symbol_font(child):
                text = text.translate(SYMBOL_FONT_MAP)
            vert = child.find("./w:rPr/w:vertAlign", ns)
            val = vert.attrib.get("{%s}val" % ns["w"]) if vert is not None else ""
            italic = child.find("./w:rPr/w:i", ns) is not None
            bold = child.find("./w:rPr/w:b", ns) is not None
            if val == "subscript":
                script_text = "<i>%s</i>" % text if italic else visible_space_text(text)
                parts.append("<sub>%s</sub>" % script_text)
            elif val == "superscript":
                script_text = "<i>%s</i>" % text if italic else visible_space_text(text)
                parts.append("<sup>%s</sup>" % script_text)
            elif italic:
                parts.append("<i>%s</i>" % text)
            elif bold:
                parts.append("<strong>%s</strong>" % text)
            elif child.find("./w:rPr/w:u", ns) is not None and not text.strip():
                parts.append("_" * max(4, len(text)))
            else:
                parts.append(visible_space_text(text))
        elif tag in ("oMath", "oMathPara"):
            math_latex = _math_latex(child).strip()
            if math_latex:
                parts.append(_legacy_latex_span(math_latex))
    text = clean_word_text(normalize_physics_markup("".join(parts)))
    text = re.sub(r"(<img\b[^>]*\/>)\s*(?=(?:[\(\uff08]\d+[\)\uff09]|[\u2460-\u2469]))", r"\n\1\n", text)
    return text


def _table_text_with_markup(table, ns, archive=None, rels=None, inline_assets=None):
    inline_assets = inline_assets if inline_assets is not None else []
    rows = []
    for tr in table.findall("./w:tr", ns):
        cells = []
        for tc in tr.findall("./w:tc", ns):
            parts = []
            for child in list(tc):
                tag = _local_name(child)
                if tag == "p":
                    text = _paragraph_text_with_markup(child, ns, archive, rels, bool(child.findall(".//o:OLEObject", ns)), inline_assets)
                    if text:
                        parts.append(text)
                elif tag == "tbl":
                    nested = _table_text_with_markup(child, ns, archive, rels, inline_assets)
                    if nested:
                        parts.append(nested)
            cells.append("<td>%s</td>" % "<br />".join(parts))
        if cells:
            rows.append("<tr>%s</tr>" % "".join(cells))
    return '<table class="question-table">%s</table>' % "".join(rows) if rows else ""


def _rels_for_document(archive):
    return _rels_for_part(archive, "word/_rels/document.xml.rels")


def _asset_from_part(archive, part_name, asset_type, rel_id=None, rel_type=None):
    if not part_name or part_name not in archive.namelist():
        return None
    data = archive.read(part_name)
    mime = mimetypes.guess_type(part_name)[0] or "application/octet-stream"
    original_file_name = os.path.basename(part_name)
    converted = _convert_windows_metafile_to_png(data, original_file_name)
    if converted:
        data = converted
        mime = "image/png"
        file_name = re.sub(r"\.(?:emf|wmf)$", ".png", original_file_name, flags=re.I)
    else:
        file_name = original_file_name
    digest = hashlib.sha256(data).hexdigest()
    asset = {
        "asset_type": asset_type,
        "file_name": file_name,
        "mime_type": mime,
        "size_bytes": len(data),
        "content_hash": digest,
        "rel_id": rel_id or "",
        "rel_type": rel_type or "",
        "source_part": part_name,
        "data_url": "data:%s;base64,%s" % (mime, base64.b64encode(data).decode("ascii")),
    }
    if converted:
        asset["original_file_name"] = original_file_name
        asset["original_mime_type"] = mimetypes.guess_type(part_name)[0] or "application/octet-stream"
    return asset


def _convert_windows_metafile_to_png(data, file_name):
    if not re.search(r"\.(?:emf|wmf)$", file_name or "", re.I):
        return None
    if os.name != "nt":
        return None

    def ps_quote(value):
        return "'" + value.replace("'", "''") + "'"

    suffix = os.path.splitext(file_name)[1].lower() or ".emf"
    src_path = ""
    out_path = ""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as src_file:
            src_file.write(data)
            src_path = src_file.name
        out_path = src_path + ".png"
        script = """
Add-Type -AssemblyName System.Drawing
$src = %s
$out = %s
$meta = [System.Drawing.Imaging.Metafile]::new($src)
$w = [Math]::Max(1, [Math]::Round($meta.Width))
$h = [Math]::Max(1, [Math]::Round($meta.Height))
$bmp = [System.Drawing.Bitmap]::new($w, $h)
$g = [System.Drawing.Graphics]::FromImage($bmp)
$g.Clear([System.Drawing.Color]::Transparent)
$g.DrawImage($meta, 0, 0, $w, $h)
$bmp.Save($out, [System.Drawing.Imaging.ImageFormat]::Png)
$g.Dispose()
$bmp.Dispose()
$meta.Dispose()
""" % (ps_quote(src_path), ps_quote(out_path))
        result = subprocess.run(
            ["powershell.exe", "-NoProfile", "-ExecutionPolicy", "Bypass", "-Command", script],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=20,
        )
        if result.returncode == 0 and os.path.exists(out_path) and os.path.getsize(out_path) > 0:
            with open(out_path, "rb") as output:
                return output.read()
    except Exception:
        return None
    finally:
        for path in (src_path, out_path):
            if path:
                try:
                    os.remove(path)
                except Exception:
                    pass
    return None


def read_docx_rich_paragraphs(file_path):
    ns = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "v": "urn:schemas-microsoft-com:vml",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "o": "urn:schemas-microsoft-com:office:office",
        "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    }
    rows = []
    try:
        with zipfile.ZipFile(file_path, "r") as archive:
            if "word/document.xml" not in archive.namelist():
                return []
            rels = _rels_for_document(archive)
            root = ET.fromstring(archive.read("word/document.xml"))
            for paragraph in root.findall(".//w:p", ns):
                assets = []
                formulas = []

                has_ole_object = bool(paragraph.findall(".//o:OLEObject", ns))
                inline_assets = []
                paragraph_text = _paragraph_text_with_markup(paragraph, ns, archive, rels, has_ole_object, inline_assets)
                assets.extend(inline_assets)

                for blip in paragraph.findall(".//a:blip", ns):
                    rid = blip.attrib.get("{%s}embed" % ns["r"]) or blip.attrib.get("{%s}link" % ns["r"])
                    rel = rels.get(rid or "")
                    asset = _asset_from_part(archive, rel.get("target") if rel else "", "formula_preview" if has_ole_object else "image", rid, rel.get("type") if rel else "")
                    if asset and asset.get("content_hash") not in {item.get("content_hash") for item in assets}:
                        asset.update(_display_size_from_container(paragraph, ns))
                        assets.append(asset)

                for image in paragraph.findall(".//v:imagedata", ns):
                    rid = image.attrib.get("{%s}id" % ns["r"])
                    rel = rels.get(rid or "")
                    asset = _asset_from_part(archive, rel.get("target") if rel else "", "formula_preview" if has_ole_object else "image", rid, rel.get("type") if rel else "")
                    if asset and asset.get("content_hash") not in {item.get("content_hash") for item in assets}:
                        asset.update(_display_size_from_container(paragraph, ns))
                        assets.append(asset)

                for obj in paragraph.findall(".//o:OLEObject", ns):
                    rid = obj.attrib.get("{%s}id" % ns["r"])
                    prog_id = obj.attrib.get("ProgID") or obj.attrib.get("Type") or "OLEObject"
                    rel = rels.get(rid or "")
                    asset = _asset_from_part(archive, rel.get("target") if rel else "", "formula_ole", rid, rel.get("type") if rel else "")
                    if asset:
                        asset["prog_id"] = prog_id
                        assets.append(asset)
                        formulas.append({
                            "format": "mathtype_ole",
                            "text": prog_id,
                            "asset_hash": asset["content_hash"],
                            "rel_id": rid or "",
                            "source_part": asset["source_part"],
                        })

                for math_node in paragraph.findall(".//m:oMath", ns) + paragraph.findall(".//m:oMathPara", ns):
                    math_text = _math_text(math_node).strip()
                    omml = _xml_bytes(math_node)
                    formulas.append({
                        "format": "omml",
                        "text": math_text,
                        "omml": omml,
                    })

                field_text = " ".join(node.text or "" for node in paragraph.findall(".//w:instrText", ns)).strip()
                if field_text:
                    formulas.append({
                        "format": "field_code",
                        "text": field_text,
                        "field_code": field_text,
                    })

                rows.append({
                    "text": paragraph_text,
                    "assets": assets,
                    "formulas": formulas,
                })
    except Exception:
        return []
    return rows


def read_docx_rich_blocks(file_path):
    ns = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "v": "urn:schemas-microsoft-com:vml",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "o": "urn:schemas-microsoft-com:office:office",
        "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    }
    try:
        with zipfile.ZipFile(file_path, "r") as archive:
            if "word/document.xml" not in archive.namelist():
                return []
            rels = _rels_for_document(archive)
            root = ET.fromstring(archive.read("word/document.xml"))
            body = root.find(".//w:body", ns)
            if body is None:
                return []
            rows = []
            for child in list(body):
                tag = _local_name(child)
                if tag == "p":
                    inline_assets = []
                    text = _paragraph_text_with_markup(child, ns, archive, rels, bool(child.findall(".//o:OLEObject", ns)), inline_assets)
                    formulas = []
                    for math_node in child.findall(".//m:oMath", ns) + child.findall(".//m:oMathPara", ns):
                        math_text = _math_text(math_node).strip()
                        formulas.append({"format": "omml", "text": math_text, "omml": _xml_bytes(math_node)})
                    rows.append({"text": text, "assets": inline_assets, "formulas": formulas, "block_type": "paragraph"})
                elif tag == "tbl":
                    inline_assets = []
                    text = _table_text_with_markup(child, ns, archive, rels, inline_assets)
                    if text:
                        rows.append({"text": text, "assets": inline_assets, "formulas": [], "block_type": "table"})
            return rows
    except Exception:
        return []


def attach_rich_content(question, rich):
    if not question or not rich:
        return
    question.setdefault("assets", [])
    question.setdefault("formulas", [])
    for asset in rich.get("assets", []):
        if asset and asset.get("content_hash") not in {item.get("content_hash") for item in question["assets"]}:
            question["assets"].append(asset)
    for formula in rich.get("formulas", []):
        if formula.get("format") == "field_code":
            continue
        if formula and formula not in question["formulas"]:
            question["formulas"].append(formula)
    question["has_image"] = any(asset.get("asset_type") == "image" for asset in question["assets"])
    question["has_formula"] = bool(question["formulas"]) or any(str(asset.get("asset_type", "")).startswith("formula_") for asset in question["assets"])


def _question_rich_text(question):
    parts = [
        question.get("stem", ""),
        question.get("content", ""),
        question.get("answer", ""),
        question.get("analysis", ""),
        question.get("explanation", ""),
    ]
    for option in question.get("options", []) or []:
        if isinstance(option, dict):
            parts.extend([option.get("content", ""), option.get("text", "")])
        else:
            parts.append(str(option))
    for sub_question in question.get("sub_questions", []) or []:
        if isinstance(sub_question, dict):
            parts.extend([sub_question.get("title", ""), sub_question.get("content", ""), sub_question.get("answer", "")])
    return "\n".join(str(part or "") for part in parts)


def attach_referenced_assets_from_rich_rows(questions, rich_rows):
    if not questions or not rich_rows:
        return
    assets_by_key = {}
    for row in rich_rows:
        for asset in row.get("assets", []) or []:
            if not asset:
                continue
            for key in (asset.get("content_hash"), asset.get("id"), asset.get("file_name")):
                if key:
                    assets_by_key[str(key)] = asset
    if not assets_by_key:
        return
    for question in questions:
        text = _question_rich_text(question)
        refs = set(re.findall(r"question-asset://([A-Za-z0-9_-]+)", text))
        for img_match in re.finditer(r"<img\b[^>]*\b(?:alt|src)=[\"']([^\"']+)[\"'][^>]*>", text, flags=re.I):
            value = img_match.group(1)
            if value.startswith("question-asset://"):
                refs.add(value.split("question-asset://", 1)[1])
            elif value:
                refs.add(value)
        if not refs:
            continue
        question.setdefault("assets", [])
        existing = {item.get("content_hash") for item in question["assets"] if item}
        for ref in refs:
            asset = assets_by_key.get(ref)
            if asset and asset.get("content_hash") not in existing:
                question["assets"].append(asset)
                existing.add(asset.get("content_hash"))
        question["has_image"] = any(asset.get("asset_type") == "image" for asset in question["assets"])


def extract_numbered_items_from_xml(file_path, definitions, comments, counters):
    namespace = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "v": "urn:schemas-microsoft-com:vml",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "o": "urn:schemas-microsoft-com:office:office",
        "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    }
    try:
        with zipfile.ZipFile(file_path, "r") as archive:
            if "word/document.xml" not in archive.namelist():
                return []
            rels = _rels_for_document(archive)
            root = ET.fromstring(archive.read("word/document.xml"))
            style_names = read_style_names_from_archive(archive)
            items = []
            body = root.find(".//w:body", namespace)
            if body is None:
                return []
            for child in list(body):
                if _local_name(child) == "tbl":
                    inline_assets = []
                    text = _table_text_with_markup(child, namespace, archive, rels, inline_assets)
                    if text:
                        items.append({
                            "text": text,
                            "number_label": "",
                            "number_kind": "",
                            "comments": [],
                            "rich": {"text": text, "assets": inline_assets, "formulas": []},
                            "is_heading": False,
                        })
                    continue
                if _local_name(child) != "p":
                    continue
                paragraph = child
                inline_assets = []
                text = _paragraph_text_with_markup(paragraph, namespace, archive, rels, bool(paragraph.findall(".//o:OLEObject", namespace)), inline_assets)
                rich = {"text": text, "assets": inline_assets, "formulas": []}
                refs = []
                for node in paragraph.findall(".//w:commentReference", namespace) + paragraph.findall(".//w:commentRangeStart", namespace) + paragraph.findall(".//w:commentRangeEnd", namespace):
                    comment_id = node.attrib.get(f"{{{namespace['w']}}}id")
                    if comment_id and comment_id not in refs:
                        refs.append(comment_id)
                num_pr = paragraph.find("./w:pPr/w:numPr", namespace)
                number_label = ""
                number_kind = ""
                if num_pr is not None:
                    num_id_node = num_pr.find("w:numId", namespace)
                    ilvl_node = num_pr.find("w:ilvl", namespace)
                    num_id = num_id_node.attrib.get(f"{{{namespace['w']}}}val") if num_id_node is not None else None
                    ilvl = ilvl_node.attrib.get(f"{{{namespace['w']}}}val") if ilvl_node is not None else "0"
                    level = definitions.get(num_id, {}).get(ilvl, {}) if num_id else {}
                    lvl_text = level.get("text", "")
                    if level.get("format") == "decimal" and "%1" in lvl_text and ilvl == "0":
                        key = (num_id, ilvl)
                        counters[key] = counters.get(key, level.get("start", 1) - 1) + 1
                        number_label = lvl_text.replace("%1", str(counters[key]))
                        if lvl_text.startswith("例"):
                            number_kind = "example"
                        elif lvl_text in ("%1", "%1.", "%1．"):
                            number_kind = "practice"
                p_style = paragraph.find("./w:pPr/w:pStyle", namespace)
                style_val = p_style.attrib.get(f"{{{namespace['w']}}}val", "") if p_style is not None else ""
                style_name = style_names.get(style_val, "")
                paragraph_comments = [comments[comment_id] for comment_id in refs if comments.get(comment_id)]
                items.append({
                    "text": text,
                    "number_label": number_label,
                    "number_kind": number_kind,
                    "comments": paragraph_comments,
                    "rich": rich,
                    "is_heading": bool(re.search(r"heading|title|\u6807\u9898", "%s %s" % (style_val, style_name), re.I)),
                })
            return items
    except Exception:
        return []


def extract_option(text):
    match = OPTION_RE.match(text)
    if match:
        return match.group(1).upper(), match.group(2).strip()
    return None, text


def extract_sub_question(text):
    match = SUB_QUESTION_RE.match(text)
    if match:
        if match.group(1):
            return "(%s)" % int(match.group(1)), match.group(3).strip()
        circled = match.group(2)
        return circled, match.group(3).strip()
    return None, text


def extract_meta(text):
    meta = {}
    year = re.search(r"(?:19|20)\d{2}", text)
    if year:
        meta["year"] = year.group(0)
    for grade in ["高一", "高二", "高三", "初一", "初二", "初三"]:
        if grade in text:
            meta["grade"] = grade
            break
    for region in ["全国", "北京", "上海", "天津", "浙江", "江苏", "山东", "广东", "湖北", "湖南", "四川"]:
        if region in text:
            meta["region"] = region
            break
    if "高考" in text:
        meta["exam_type"] = "高考真题"
    elif "模拟" in text or "联考" in text or "适应性" in text:
        meta["exam_type"] = "模拟题"
    elif "期中" in text:
        meta["exam_type"] = "期中考试"
    elif "期末" in text:
        meta["exam_type"] = "期末考试"
    elif "月考" in text:
        meta["exam_type"] = "月考"
    return meta


def extract_types_from_text(text, options=None):
    mapping = {
        "多选": "multi",
        "单选": "single",
        "选择": "single",
        "填空": "fill",
        "实验": "experiment",
        "计算": "calculation",
        "解答": "problem",
        "简答": "short",
        "作图": "drawing",
        "判断": "judge",
    }
    types = [value for key, value in mapping.items() if key in text]
    if not types and options:
        types = ["single"]
    return types or ["fill"]


def answer_letters(answer):
    text = re.sub(r"【.*?】", "", str(answer or "")).upper()
    return re.findall(r"[A-G]", text)


def classify_question(stem, options=None, answer=""):
    letters = answer_letters(answer)
    if options or re.search(r"[（(]\s*[　\s]{2,}[）)]", stem or ""):
        return ["multi" if len(set(letters)) >= 2 else "single"]
    if re.search(r"实验", stem or ""):
        return ["experiment"]
    if re.search(r"解答|计算|综合", stem or ""):
        return ["problem"]
    return extract_types_from_text(stem or "", options)


def split_image_only_options(question):
    options = question.get("options") or []
    if len(options) != 2:
        return
    labels = [str(option.get("label", "")).upper() for option in options]
    if labels != ["A", "C"]:
        return
    next_options = []
    target_labels = [["A", "B"], ["C", "D"]]
    for option, pair_labels in zip(options, target_labels):
        content = option.get("content", "")
        images = re.findall(r"<img\b[^>]*>", content, flags=re.I)
        remainder = re.sub(r"<img\b[^>]*>", "", content, flags=re.I)
        remainder = re.sub(r"[A-G][\.\u3001\uff0e\uff0c\uff1a:;；\s]*", "", remainder, flags=re.I).strip()
        if len(images) != 2 or remainder:
            return
        for label, image in zip(pair_labels, images):
            next_options.append({"label": label, "content": image, "is_correct": False})
    question["options"] = next_options


def split_packed_options(question):
    options = question.get("options") or []
    next_options = []
    changed = False
    for option in options:
        raw = "%s. %s" % (option.get("label", "A"), option.get("content", ""))
        matches = list(re.finditer(r"(^|[\r\n\t\f])\s*([A-G])[\.\uff0e]\s*", raw, flags=re.I))
        if len(matches) < 2:
            next_options.append(option)
            continue
        expanded = []
        for index, match in enumerate(matches):
            prefix = match.group(1) or ""
            content_start = match.end()
            content_end = (matches[index + 1].start() + len(matches[index + 1].group(1) or "")) if index + 1 < len(matches) else len(raw)
            content = raw[content_start:content_end].strip()
            if content:
                expanded.append({"label": match.group(2).upper(), "content": content, "is_correct": False})
        if len(expanded) >= 2:
            next_options.extend(expanded)
            changed = True
        else:
            next_options.append(option)
    if changed:
        question["options"] = next_options


def normalize_subquestion_image_positions(question):
    stem = question.get("stem", "")
    if not stem or "<img" not in stem:
        return
    sub_re = r"(?:[\(\uff08]\d+[\)\uff09]|[\u2460-\u2469])"
    first_sub = re.search(r"(^|\n)\s*%s" % sub_re, stem)
    if not first_sub:
        return
    moved = []

    def pull_image(match):
        moved.append(match.group(2))
        return "\n" + match.group(3)

    next_stem = re.sub(
        r"(^|\n)\s*(<img\b[^>]*\/>)\s*\n\s*(%s)" % sub_re,
        pull_image,
        stem,
    )
    if not moved:
        return
    first_sub = re.search(r"(^|\n)\s*%s" % sub_re, next_stem)
    if not first_sub:
        question["stem"] = next_stem.strip()
        return
    insert_at = first_sub.start()
    if insert_at == 0:
        line_end = next_stem.find("\n", first_sub.end())
        insert_at = len(next_stem) if line_end < 0 else line_end
    image_block = "\n".join(moved)
    question["stem"] = (next_stem[:insert_at].rstrip() + "\n" + image_block + "\n" + next_stem[insert_at:].lstrip()).strip()


def normalize_leading_image_positions(question):
    stem = question.get("stem", "")
    if not stem or "<img" not in stem:
        return
    inline_match = re.match(r"^\s*((?:<img\b[^>]*>\s*)+)([\s\S]*)$", stem, flags=re.I)
    if inline_match and inline_match.group(2).strip():
        stem = inline_match.group(1).strip() + "\n" + inline_match.group(2).lstrip()
    lines = stem.splitlines()
    leading = []
    while lines and re.match(r"^\s*(?:<img\b[^>]*>\s*)+\s*$", lines[0], flags=re.I):
        leading.append(lines.pop(0).strip())
    if not leading:
        return
    boundary_re = r"^\s*(?:[A-G][\.\uff0e]|[\(\uff08]\d+[\)\uff09]|[\u2460-\u2469])"
    insert_at = next((index for index, line in enumerate(lines) if re.match(boundary_re, line, flags=re.I)), min(len(lines), 1))
    if insert_at == 0 and lines:
        insert_at = 1
    lines[insert_at:insert_at] = leading
    question["stem"] = "\n".join(lines).strip()


def finalize_question_type(question):
    split_packed_options(question)
    normalize_leading_image_positions(question)
    normalize_subquestion_image_positions(question)
    question["question_types"] = classify_question(question.get("stem", ""), question.get("options", []), question.get("answer", ""))
    return question


def derive_topic_from_filename(file_path):
    base = os.path.splitext(os.path.basename(file_path))[0]
    patterns = [
        r"专题\d+[-：:](.+)$",
        r"实验专题\d+[-：:](.+)$",
        r"解答题专题\d+[-：:](.+)$",
    ]
    for pattern in patterns:
        match = re.search(pattern, base)
        if match:
            return match.group(1).strip()
    return ""


def is_topic_heading(text):
    if not text:
        return False
    generic = ["选择题", "非选择题", "填空题", "实验题", "计算题", "解答题", "答案", "解析", "参考答案"]
    if any(word in text for word in generic):
        return False
    return bool(re.match(r"^(专题\d+|实验专题\d+|解答题专题\d+|[一二三四五六七八九十]+[、.．])", text))


def normalize_heading_text(text):
    text = re.sub(r"<[^>]+>", "", text or "")
    text = text.replace("\u3000", " ")
    text = re.sub(r"\s+", "", text)
    text = re.sub(r"[:：]+$", "", text)
    return text

def is_pure_question_type_title(text):
    text = normalize_heading_text(text)
    return bool(re.match(r"^(?:\d+[\u3001\uff0e\.、]?|[\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341]+[\u3001\uff0e\.、]?)?(\u5355\u9009\u9898|\u591a\u9009\u9898|\u9009\u62e9\u9898|\u5b9e\u9a8c\u9898|\u89e3\u7b54\u9898|\u7efc\u5408\u9898|\u586b\u7a7a\u9898)$", text))


def is_exam_section_heading(text):
    text = normalize_heading_text(text)
    return is_pure_question_type_title(text) or bool(EXAM_SECTION_RE.match(text)) or bool(re.match(r"^\d+[\u3001\uff0e\.\s]*(\u5355\u9009\u9898|\u591a\u9009\u9898|\u9009\u62e9\u9898|\u5b9e\u9a8c\u9898|\u89e3\u7b54\u9898|\u7efc\u5408\u9898|\u586b\u7a7a\u9898)$", text))


def is_section_heading_like(text):
    text = normalize_heading_text(text)
    return bool(re.match(r"^(?:\d+|[\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341]+)[\u3001\uff0e\.]?(?:\u5355\u9009\u9898|\u591a\u9009\u9898|\u9009\u62e9\u9898|\u5b9e\u9a8c\u9898|\u89e3\u7b54\u9898|\u7efc\u5408\u9898|\u586b\u7a7a\u9898)(?:[\uff08(].*)?$", text))


def should_inline_option(question):
    if not question:
        return False
    text = "".join([
        str(question.get("stem", "")),
        str(question.get("knowledge_point", "")),
        str(question.get("section_title", "")),
    ])
    return bool(question.get("sub_questions")) or any(token in text for token in ("\u5b9e\u9a8c", "\u89e3\u7b54", "\u7efc\u5408"))


def clean_topic_heading(text):
    text = re.sub(r"^(专题\d+|实验专题\d+|解答题专题\d+)[-：:、.\s]*", "", text).strip()
    text = re.sub(r"^[一二三四五六七八九十]+[、.．]\s*", "", text).strip()
    return text[:80]


def new_question(stem, index=None, knowledge_point=None):
    knowledge_points = [knowledge_point] if knowledge_point else []
    return {
        "id": str(uuid.uuid4()),
        "index": index,
        "stem": stem.strip(),
        "options": [],
        "sub_questions": [],
        "answer": "",
        "analysis": "",
        "knowledge_point": knowledge_point or "",
        "knowledge_points": knowledge_points,
        "question_types": extract_types_from_text(stem),
        "assets": [],
        "formulas": [],
        "has_image": False,
        "has_formula": False,
    }


def append_text(question, text):
    question["stem"] = (question["stem"] + "\n" + text).strip()
    if question["sub_questions"]:
        question["sub_questions"][-1]["content"] += "\n" + text


def parse_question_block(paragraphs, default_topic=None):
    questions = []
    current = None
    current_topic = default_topic or ""
    for paragraph in paragraphs:
        text = paragraph.strip()
        if not text:
            continue
        if is_section_heading_like(text) or is_pure_question_type_title(text):
            continue
        if is_topic_heading(text):
            topic = clean_topic_heading(text)
            if topic:
                current_topic = topic
            continue
        number, content = extract_question_number(text)
        if number is not None:
            if is_pure_question_type_title(content):
                continue
            if current:
                questions.append(finalize_question_type(current))
            current = new_question(content, len(questions), current_topic)
            continue
        if not current:
            continue
        option, option_content = extract_option(text)
        if option:
            if should_inline_option(current):
                append_text(current, text)
            else:
                current["options"].append({"label": option, "content": option_content, "is_correct": False})
            continue
        sub_label, sub_content = extract_sub_question(text)
        if sub_label:
            current["sub_questions"].append({"title": sub_label, "content": sub_content, "answer": ""})
            current["stem"] = (current["stem"] + "\n" + f"{sub_label} {sub_content}").strip()
            continue
        append_text(current, text)
    if current:
        questions.append(finalize_question_type(current))
    return questions


def split_answer_section(paragraphs):
    markers = ("\u53c2\u8003\u7b54\u6848", "\u7b54\u6848\u4e0e\u89e3\u6790", "\u7b54\u6848\u53ca\u89e3\u6790")
    for index, paragraph in enumerate(paragraphs):
        text = paragraph.strip()
        if ANSWER_TITLE_RE.search(text) or any(marker in text for marker in markers):
            return paragraphs[:index], paragraphs[index + 1 :]
    return paragraphs, []

def parse_answers(paragraphs):
    answers = []
    current = None
    for paragraph in paragraphs:
        text = paragraph.strip()
        if not text:
            continue
        number, content = extract_question_number(text)
        if number is not None:
            if current:
                answers.append(current)
            current = {"number": number, "answer": content, "explanation": ""}
            continue
        if current:
            if current["explanation"]:
                current["explanation"] += "\n"
            current["explanation"] += text
    if current:
        answers.append(current)
    return answers


def parse_exam_answer_blocks(paragraphs):
    answers = []
    current = None
    mode = "answer"
    for paragraph in paragraphs:
        text = paragraph.strip()
        if not text:
            continue
        if ANSWER_TITLE_RE.match(text):
            continue
        number, content = extract_question_number(text)
        if number is not None:
            if current:
                answers.append(current)
            current = {"number": number, "answer": "", "explanation": ""}
            mode = "answer"
            marker = ANALYSIS_MARK_RE.search(content)
            if marker:
                current["answer"] = content[: marker.start()].strip()
                current["explanation"] = content[marker.end() :].strip()
                mode = "analysis"
            else:
                current["answer"] = content.strip()
            continue
        if not current:
            continue
        marker = ANALYSIS_MARK_RE.search(text)
        if marker:
            before = text[: marker.start()].strip()
            after = text[marker.end() :].strip()
            if before:
                current["answer"] = (current["answer"] + "\n" + before).strip()
            if after:
                current["explanation"] = (current["explanation"] + "\n" + after).strip()
            mode = "analysis"
            continue
        if mode == "answer":
            current["answer"] = (current["answer"] + "\n" + text).strip()
        else:
            current["explanation"] = (current["explanation"] + "\n" + text).strip()
    if current:
        answers.append(current)
    return answers


def parse_exam_question_block(paragraphs):
    questions = []
    current = None
    expected_number = None
    current_section = ""
    for paragraph in paragraphs:
        text = paragraph.strip()
        if not text:
            continue
        if is_exam_section_heading(text):
            if current:
                questions.append(finalize_question_type(current))
                current = None
            expected_number = None
            current_section = text
            continue
        number, content = extract_question_number(text)
        if number is not None and (expected_number is None or number == expected_number):
            if is_pure_question_type_title(content):
                current_section = content
                expected_number = None
                continue
            if current:
                questions.append(finalize_question_type(current))
            current = new_question(content, len(questions), "")
            current["number"] = number
            current["section_title"] = current_section
            expected_number = number + 1
            continue
        if not current:
            continue
        option, option_content = extract_option(text)
        if option:
            if any(word in current_section for word in ("\u5b9e\u9a8c\u9898", "\u89e3\u7b54\u9898", "\u7efc\u5408\u9898")) or current["sub_questions"]:
                append_text(current, text)
            else:
                current["options"].append({"label": option, "content": option_content, "is_correct": False})
            continue
        sub_label, sub_content = extract_sub_question(text)
        if sub_label:
            current["sub_questions"].append({"title": sub_label, "content": sub_content, "answer": ""})
            current["stem"] = (current["stem"] + "\n" + f"{sub_label} {sub_content}").strip()
            continue
        append_text(current, text)
    if current:
        questions.append(finalize_question_type(current))
    return questions


def parse_lecture_questions(paragraphs, default_topic=None):
    return parse_question_block(paragraphs, default_topic)


def parse_lecture_numbered_items(items, default_topic=None):
    questions = []
    current = None
    current_topic = default_topic or ""

    def finish_current():
        nonlocal current
        if current:
            comments = [comment for comment in current.pop("_comments", []) if comment]
            if comments and not current.get("answer"):
                current["answer"] = "\n".join(comments)
            questions.append(finalize_question_type(current))
            current = None

    for item in items:
        text = item.get("text", "").strip()
        rich = item.get("rich", {})
        number_kind = item.get("number_kind", "")
        is_heading = bool(item.get("is_heading"))
        if not text and not number_kind:
            if current and (rich.get("assets") or rich.get("formulas")):
                attach_rich_content(current, rich)
            continue
        if is_heading and number_kind != "example":
            finish_current()
            topic = clean_topic_heading(text)
            if topic:
                current_topic = topic
            continue
        if text and (is_section_heading_like(text) or is_pure_question_type_title(text)):
            finish_current()
            continue
        if text and is_topic_heading(text):
            finish_current()
            topic = clean_topic_heading(text)
            if topic:
                current_topic = topic
            continue
        if number_kind in ("example", "practice"):
            finish_current()
            stem = text.strip()
            current = new_question(stem, len(questions), current_topic)
            current["source_type"] = number_kind
            current["_comments"] = list(item.get("comments", []))
            attach_rich_content(current, rich)
            continue
        if not current or not text:
            if current:
                current.setdefault("_comments", []).extend(item.get("comments", []))
                attach_rich_content(current, rich)
            continue
        current.setdefault("_comments", []).extend(item.get("comments", []))
        attach_rich_content(current, rich)
        option, option_content = extract_option(text)
        if option:
            if should_inline_option(current):
                append_text(current, text)
            else:
                current["options"].append({"label": option, "content": option_content, "is_correct": False})
            continue
        sub_label, sub_content = extract_sub_question(text)
        if sub_label:
            current["sub_questions"].append({"title": sub_label, "content": sub_content, "answer": ""})
            current["stem"] = (current["stem"] + "\n" + f"{sub_label} {sub_content}").strip()
            continue
        append_text(current, text)

    finish_current()
    return questions


def parse_exam_questions(paragraphs, default_topic=None, rich_rows=None):
    question_part, answer_part = split_answer_section(paragraphs)
    question_rich = rich_rows[: len(question_part)] if rich_rows else None
    questions = parse_exam_question_block(question_part)
    if question_rich:
        current_index = -1
        expected_number = None
        for paragraph_index, text in enumerate(question_part):
            rich = question_rich[paragraph_index] if paragraph_index < len(question_rich) else {}
            stripped = (text or "").strip()
            number, _content = extract_question_number(stripped)
            if number is not None and (expected_number is None or number == expected_number):
                current_index += 1
                expected_number = number + 1
            if 0 <= current_index < len(questions) and (rich.get("assets") or rich.get("formulas")):
                attach_rich_content(questions[current_index], rich)
    answers = parse_exam_answer_blocks(answer_part)
    by_index = {idx: answer for idx, answer in enumerate(answers)}
    by_number = {answer["number"]: answer for answer in answers}
    for idx, question in enumerate(questions):
        answer = by_number.get(question.get("number") or idx + 1) or by_index.get(idx)
        if answer:
            question["answer"] = answer.get("answer", "")
            question["analysis"] = answer.get("explanation", "")
            finalize_question_type(question)
    return questions


def extract_comments(file_path):
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    try:
        with zipfile.ZipFile(file_path, "r") as archive:
            if "word/comments.xml" not in archive.namelist():
                return []
            root = ET.fromstring(archive.read("word/comments.xml"))
            comments = []
            for comment in root.findall(".//w:comment", namespace):
                texts = [node.text or "" for node in comment.findall(".//w:t", namespace)]
                comments.append("".join(texts).strip())
            return comments
    except Exception:
        return []


def extract_paragraphs_from_docx(file_path):
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    paragraphs = []
    with zipfile.ZipFile(file_path, "r") as archive:
        if "word/document.xml" not in archive.namelist():
            raise ValueError("word/document.xml not found")
        root = ET.fromstring(archive.read("word/document.xml"))
        for paragraph in root.findall(".//w:p", namespace):
            text = _paragraph_text_with_markup(paragraph, {
                "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
                "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
            })
            if text:
                paragraphs.append(text)
    return paragraphs


def extract_topics(paragraphs):
    topics = []
    for paragraph in paragraphs:
        text = paragraph.strip()
        if text and is_topic_heading(text):
            topic = clean_topic_heading(text)
            if topic and topic not in topics:
                topics.append(topic)
    return topics


def quality_report(questions):
    warnings = {}
    for question in questions:
        if not question.get("stem"):
            warnings["missing_stem"] = warnings.get("missing_stem", 0) + 1
        if not question.get("answer"):
            warnings["missing_answer"] = warnings.get("missing_answer", 0) + 1
        if question.get("options") and len(question["options"]) < 2:
            warnings["few_options"] = warnings.get("few_options", 0) + 1
    return {"warnings": warnings, "parsed_items": len(questions)}


def main():
    if len(sys.argv) < 3:
        print(json.dumps({"error": "usage: parse_word.py <file_path> <source_type> [knowledge_tree]"}, ensure_ascii=False))
        sys.exit(1)

    file_path = sys.argv[1]
    source_type = sys.argv[2]
    if not os.path.exists(file_path):
        print(json.dumps({"error": f"file not found: {file_path}"}, ensure_ascii=False))
        sys.exit(1)

    try:
        if os.environ.get("GEWU_FORCE_DOCX_XML_FALLBACK") == "1":
            raise ImportError("forced docx xml fallback")
        from docx import Document
        doc = Document(file_path)
        rich_rows = read_docx_rich_blocks(file_path) or read_docx_rich_paragraphs(file_path)
        comment_assets = read_docx_comment_assets(file_path)
        paragraphs = [row.get("text", "") for row in rich_rows] or [paragraph.text for paragraph in doc.paragraphs]
        numbered_items = extract_numbered_items(doc, file_path)
    except ImportError:
        try:
            paragraphs = extract_paragraphs_from_docx(file_path)
            rich_rows = read_docx_rich_blocks(file_path) or read_docx_rich_paragraphs(file_path)
            comment_assets = read_docx_comment_assets(file_path)
            numbered_items = []
        except Exception as exc:
            print(json.dumps({"error": f"cannot open Word document without python-docx: {exc}"}, ensure_ascii=False))
            sys.exit(1)
    except Exception as exc:
        try:
            paragraphs = extract_paragraphs_from_docx(file_path)
            rich_rows = read_docx_rich_blocks(file_path) or read_docx_rich_paragraphs(file_path)
            comment_assets = read_docx_comment_assets(file_path)
            numbered_items = []
        except Exception:
            print(json.dumps({"error": f"cannot open Word document: {exc}"}, ensure_ascii=False))
            sys.exit(1)
    if source_type == "auto":
        source_type = "exam" if any("参考答案" in paragraph or "答案" in paragraph for paragraph in paragraphs) else "lecture"

    default_topic = derive_topic_from_filename(file_path) if source_type == "lecture" else ""
    if source_type == "exam":
        questions = parse_exam_questions(paragraphs, default_topic, rich_rows)
    elif numbered_items:
        questions = parse_lecture_numbered_items(numbered_items, default_topic)
    else:
        questions = []
    asset_rows = list(rich_rows or [])
    if comment_assets:
        asset_rows.append({"assets": comment_assets})
    attach_referenced_assets_from_rich_rows(questions, asset_rows)

    result = {
        "success": True,
        "source_type": source_type,
        "count": len(questions),
        "questions": questions,
        "topics": extract_topics(paragraphs) or ([default_topic] if default_topic else []),
        "knowledge_points": sorted({kp for question in questions for kp in question.get("knowledge_points", []) if kp}),
        "quality_report": quality_report(questions),
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
